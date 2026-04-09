"""Generate SunghunKim_Milestone2.docx"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ── Styles ──────────────────────────────────────────────────────────────────
style_normal = doc.styles["Normal"]
style_normal.font.name = "Calibri"
style_normal.font.size = Pt(11)

def set_heading(paragraph, level=1):
    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
    run.bold = True
    if level == 1:
        run.font.size = Pt(14)
    elif level == 2:
        run.font.size = Pt(12)

def heading(doc, text, level=1):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14) if level == 1 else Pt(12)
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    return p

def body(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    return p

def bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    p.paragraph_format.space_after = Pt(3)
    return p

def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        hdr[i].paragraphs[0].runs[0].bold = True
    for row_data in rows:
        row = table.add_row().cells
        for i, val in enumerate(row_data):
            row[i].text = val
            row[i].paragraphs[0].runs[0].font.size = Pt(10)
    doc.add_paragraph()


# ── Title ────────────────────────────────────────────────────────────────────
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("Mapping the U.S. Digital Economy")
run.bold = True
run.font.size = Pt(18)

doc.add_paragraph()
info = doc.add_paragraph()
info.add_run("Student: ").bold = True
info.add_run("Sunghun Kim\n")
info.add_run("Advisor: ").bold = True
info.add_run("Fernando Santos\n")
info.add_run("Milestone: ").bold = True
info.add_run("2")

doc.add_paragraph()

# ── Vision Statement ────────────────────────────────────────────────────────
heading(doc, "Vision Statement")
body(doc,
    "The U.S. digital economy encompasses thousands of technology companies across diverse "
    "sectors, yet no comprehensive, publicly accessible map exists that connects what "
    "technologies developers are adopting, what companies are investing in R&D, and where "
    "financial returns are materializing. This project builds a systematic map of 4,001 U.S. "
    "tech companies across 19 major tech hubs, classifying each company by sector, revenue "
    "model, and work regime using a taxonomy developed by Prof. Santos. By integrating "
    "developer survey data (Stack Overflow, 2017–2025), public financial filings (SEC EDGAR), "
    "and labor market statistics (BLS), the project creates an interactive investment "
    "intelligence dashboard that gives investors, students, and policymakers actionable "
    "insights into which sectors of the U.S. digital economy are growing, innovating, and "
    "worth watching. The central question driving this work is: where should the next "
    "investment dollar in the U.S. digital economy go?"
)

# ── Background ───────────────────────────────────────────────────────────────
heading(doc, "Background")
body(doc,
    "Before starting this project, I was encouraged to join a team or continue someone else's "
    "ongoing senior project. However, I wanted to create something independently in an area of "
    "personal interest that would also develop practical skills useful for my professional career."
)
body(doc,
    "While discussing this goal with Professor Santos, he shared an outline of a grant proposal "
    "he had written. He suggested I could start working on a project inspired by his idea, but "
    "with directions changed to match my own interests. After reading through his project "
    "outline, I was fascinated by the approach — particularly because the core of the work "
    "involved data engineering, which I had not practiced before. This technical challenge and "
    "practical impact became the motivation for choosing to do this project."
)
body(doc,
    "The original scope was limited to Michigan-based tech companies (~200 companies). However, "
    "as data collection progressed, it became clear that restricting the dataset to a single "
    "state created significant data limitations: fewer companies meant narrower sector "
    "representation, weaker statistical signals, and limited room for geographic comparison. "
    "Expanding the scope to the entire United States (4,001 companies across 19 major tech "
    "hubs) resolved these limitations and opened up substantially more analytical directions — "
    "including state-level comparisons, hub specialization analysis, and a more robust "
    "Opportunity Score backed by a larger and more diverse dataset."
)
body(doc,
    "Several existing resources informed the project's design. The U.S. Bureau of Economic "
    "Analysis publishes an annual Digital Economy report that estimates the digital economy's "
    "share of U.S. GDP [1], but it relies on NAICS industry codes that pre-date AI-era sectors "
    "and cannot distinguish categories such as 'AI foundation models' from broader software "
    "services. The Stack Overflow Annual Developer Survey [2] is the largest and most widely "
    "cited source of data on developer tools and practices, covering 65,000+ respondents "
    "globally each year; this project uses the 2017–2025 editions as the Learning layer of "
    "the analysis. SEC EDGAR [3] is the U.S. Securities and Exchange Commission's public "
    "database of company financial filings, used here to extract revenue, R&D, and net income "
    "for publicly traded companies. The U.S. Bureau of Labor Statistics Quarterly Census of "
    "Employment and Wages [4] provides sector-level employment and wage data used in the "
    "Investing layer. builtin.com [5] is a tech industry job and company directory covering "
    "major U.S. tech hubs; it serves as the primary source for the company dataset in this project."
)

# ── Normative and Ethical Considerations ─────────────────────────────────────
heading(doc, "Normative and Ethical Considerations")

heading(doc, "1. Privacy", level=2)
body(doc, "Scraping company data without permission could violate individual and organizational privacy rights.")
bullet(doc, "Only publicly available data was collected from builtin.com's public company directory.")
bullet(doc, "No individual employee data is collected; all analysis is at the company or sector level.")
bullet(doc, "Challenge: Distinguishing truly public data from data that is technically accessible but not intended for bulk collection.")

heading(doc, "2. Transparency", level=2)
body(doc, "If data sources are not transparent, the reliability of the analysis cannot be assessed, leading to misinformed decisions.")
bullet(doc, "All four data sources are documented in a dedicated 'Data & Methods' section of the dashboard.")
bullet(doc, "Data limitations are explicitly disclosed with warning banners (e.g., SEC EDGAR covers public companies only).")
bullet(doc, "Full pipeline code is published on GitHub (github.com/Huni-code/Senior-Project_SunghunK).")

heading(doc, "3. Fairness", level=2)
body(doc, "The analysis could misrepresent smaller companies if data is skewed toward large, publicly traded firms.")
bullet(doc, "The dataset includes 896 startups (< 50 employees) alongside mid-size and enterprise companies.")
bullet(doc, "SEC EDGAR coverage limitation (public companies only, 73% of dataset) is explicitly flagged with warning banners in Sections 3 and 4 of the dashboard.")
bullet(doc, "Challenge: Private startups are systematically underrepresented in financial analysis due to limited public disclosure requirements.")

heading(doc, "4. AI Classification Bias (new since M1)", level=2)
body(doc, "LLM-based classification may systematically misclassify companies in ambiguous sectors.")
bullet(doc, "Classification uses a structured prompt with Prof. Santos's exact taxonomy definitions.")
bullet(doc, "Sector distributions were spot-checked against known companies (e.g., Google → Advertising & attention, Stripe → Fintech & payments).")

# ── Success Criteria ─────────────────────────────────────────────────────────
heading(doc, "Product Success Criteria")
add_table(doc,
    ["Criterion", "Minimum", "Excellence", "Current Status"],
    [
        ["Data Coverage", "1000+ companies, 5+ sectors", "2000+ companies, 10+ sectors", "4,001 companies, 16 sectors"],
        ["Geographic Scope", "Michigan only", "Multi-state", "50 states, 19 tech hubs"],
        ["Analytical Depth", "Basic visualizations", "Predictive model, network graphs", "Regression + Opportunity Score"],
        ["Data Sources", "1 source", "3+ sources", "4 sources integrated"],
        ["Dashboard", "Static report", "Interactive, deployed", "Deployed on Streamlit Cloud"],
        ["Classification", "Manual tagging", "LLM-assisted", "Claude Haiku, 4,001 companies"],
        ["Report", "15+ pages", "25+ pages", "In progress"],
        ["Work Regime", "Not planned", "All 9 categories labeled", "Planned"],
    ]
)

# ── Approach and Implementation / Results ─────────────────────────────────────
heading(doc, "Approach and Implementation / Results")

heading(doc, "Visible Accomplishment", level=2)
body(doc,
    "A fully deployed, interactive investment intelligence dashboard is live on Streamlit Cloud "
    "(github.com/Huni-code/Senior-Project_SunghunK), backed by a SQLite database containing "
    "4,001 classified U.S. tech companies and data from four integrated sources. The dashboard "
    "allows users to adjust investment weights in real-time and receive updated sector rankings, "
    "company drilldowns, and revenue forecasts through 2027."
)

heading(doc, "Data Pipeline", level=2)
bullet(doc,
    " A custom Playwright-based web scraper collected 7,601 raw entries across 19 U.S. "
    "tech hubs (Boston, SF, NYC, Chicago, Austin, Seattle, Denver, Atlanta, LA, Dallas, "
    "and others). After deduplication by URL, 4,001 unique companies were retained. Each "
    "company was enriched with state, employee count, and size classification "
    "(Startup / Small / Mid-size / Enterprise).",
    bold_prefix="Company Data (builtin.com):"
)
bullet(doc,
    " Each company's name, description, and sector tags were sent to Claude Haiku "
    "(claude-haiku-4-5-20251001) with a structured prompt based on Prof. Santos's taxonomy, "
    "producing sector (16 categories) and revenue model (5 categories) labels for all 4,001 "
    "companies at a cost of ~$4.",
    bold_prefix="LLM Classification:"
)
bullet(doc,
    " A matching pipeline mapped companies to SEC CIK identifiers and retrieved revenue, "
    "R&D expense, and net income for 2,937 public companies (2015–2024), producing 15,952 "
    "financial records.",
    bold_prefix="SEC EDGAR Financial Data:"
)
bullet(doc,
    " Nine years of Stack Overflow survey data were processed to extract tool adoption "
    "trends, AI tool adoption rates, developer salary trends, and developer type "
    "distributions. Only professional developers were included (MainBranch filter).",
    bold_prefix="Stack Overflow Developer Survey (2017–2025):"
)
bullet(doc,
    " Employment and wage trends were collected from the U.S. Bureau of Labor Statistics "
    "for 11 tech-related industry categories (2015–2025). Since BLS uses its own industry "
    "classification system (NAICS), only categories that clearly matched the project's sectors "
    "were included — for example, 'Data Processing & Hosting' maps well to cloud/infrastructure "
    "sectors, while newer AI-specific sectors have no direct BLS equivalent and were left out.",
    bold_prefix="BLS Employment Data:"
)

heading(doc, "System Architecture", level=2)
body(doc, "Figure 1 illustrates the end-to-end data pipeline, from raw data sources through processing and storage to the deployed dashboard.")
doc.add_picture("architecture_diagram.png", width=Inches(6.0))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
p = doc.add_paragraph("Figure 1. System architecture and data pipeline.")
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.runs[0].italic = True
p.runs[0].font.size = Pt(9)

heading(doc, "Database Model", level=2)
body(doc, "The SQLite database (data/companies.db) is organized into four layers:")
add_table(doc,
    ["Table", "Rows", "Key Columns", "Purpose"],
    [
        ["companies_deduped", "4,001", "name, hub, state, company_size, employees_count", "Core company registry"],
        ["company_classifications", "4,001", "company_id, sector, revenue_model", "LLM taxonomy labels"],
        ["sec_financials", "15,952", "company_id, year, revenue, rd_expense, net_income", "Financial data (2015–2024)"],
        ["sec_cik_map", "2,937", "company_id, cik", "SEC identifier mapping"],
        ["bls_employment", "110", "sector, year, employees, avg_hourly_wage", "Labor market trends"],
        ["so_tools_trend", "57", "year, tool, usage_pct", "Developer tool adoption"],
        ["so_ai_adoption", "63", "year, ai_tool, usage_pct", "AI tool adoption (2023–2025)"],
        ["so_salary_trend", "9", "year, median_salary, p25_salary, p75_salary", "Developer salary trends"],
        ["so_devtype_trend", "234", "year, dev_type, pct", "Developer role distribution"],
        ["so_desire_gap", "~120", "year, tool, have_pct, want_pct, gap", "Technology demand signals"],
    ]
)

heading(doc, "Opportunity Score", level=2)
body(doc,
    "The dashboard's core analytical output is the Opportunity Score, which integrates all "
    "three analytical layers into a single 0–100 ranking per sector:"
)
p = doc.add_paragraph()
p.paragraph_format.left_indent = Inches(0.5)
p.add_run("Opportunity Score = wL × Learning  +  wI × Inventing  +  wV × Investing").italic = True

body(doc,
    "Learning is derived from SO Survey technology adoption relevance per sector; Inventing "
    "from SEC R&D/Revenue ratio (2020–2024 average); Investing from revenue CAGR (2019–2024). "
    "Default weights are Learning=40%, Inventing=30%, Investing=30%, adjustable in real-time "
    "via a sidebar slider with preset investor profiles (Growth, Value, Balanced)."
)

heading(doc, "Key Findings So Far", level=2)

heading(doc, "Learning — Technology Adoption (SO Survey, 2023–2025)", level=2)
bullet(doc, "ChatGPT dominated AI tool adoption at 88% in 2023. However, Claude surged from 8% (2024) to 43% (2025), a +35 percentage point jump in a single year, while DeepSeek entered at 24% in 2025 with no prior presence — indicating the AI tooling landscape is rapidly shifting.")
bullet(doc, "AI/ML Engineer appeared as an independent developer role category for the first time in the 2025 survey, reflecting a structural change in how the industry defines and hires technical roles.")
bullet(doc, "The desire gap analysis shows PyTorch (+2.1pp) and TensorFlow (+2.0pp) have growing unmet demand among developers, while React (-6pp), Node.js (-8pp), and Angular (-7pp) are in saturation — developers already have these skills and are not actively seeking more.")

heading(doc, "Inventing — R&D Investment (SEC EDGAR, 2015–2024)", level=2)
bullet(doc, "Search Engines (~210%) and AI Foundation Models (~200%) have the highest R&D-to-revenue ratios across all sectors — both spend more on R&D than they generate in revenue.")
bullet(doc, "AI Foundation Models R&D spending accelerated sharply after 2022, a trend clearly visible in the R&D expense trend chart, aligning with the public release of large language models.")
bullet(doc, "Fintech and E-commerce show the lowest R&D intensity across all sectors, sitting at the bottom of the R&D/Revenue ranking.")

heading(doc, "Investing — Growth & Market Structure (SEC EDGAR + BLS)", level=2)
bullet(doc, "All 16 sectors showed positive revenue growth between 2019 and 2024. AI Foundation Models led at +386%, followed by Advertising & Attention (+374%) and E-learning (+311%).")
bullet(doc, "Enterprise/ERP/HRM is the largest sector by company count (902 companies, 22.5% of the dataset) but ranks among the lowest in R&D intensity.")
bullet(doc, "BLS employment divergence: Data Processing & Hosting grew +62.5% in employment (2015–2024), while Telecommunications declined -27% over the same period.")
bullet(doc, "Geographic concentration: CA, TX, NY, and MA account for 75% of all tech companies in the dataset.")

# ── Milestone-Specific: Updates Since M1 ────────────────────────────────────
heading(doc, "Milestone-Specific: Updates Since Milestone 1")
body(doc,
    "Since Milestone 1, the project underwent a significant scope expansion and a complete "
    "technical rebuild. The following updates were made in direct response to Professor "
    "Santos's feedback:"
)
bullet(doc,
    " The Streamlit dashboard is now deployed and publicly accessible on Streamlit Cloud. "
    "Automatic redeployment is configured via GitHub push.",
    bold_prefix="Dashboard deployment (addressed M1 feedback):"
)
bullet(doc,
    " A full database schema has been documented (see table above), including all table "
    "names, column names, row counts, and purposes — as requested by Prof. Santos.",
    bold_prefix="Database model documentation (addressed M1 feedback):"
)
bullet(doc,
    " The project expanded from Michigan (~200 companies) to the entire United States "
    "(4,001 companies across 19 tech hubs and all 50 states).",
    bold_prefix="Scope expansion:"
)
bullet(doc,
    " Prof. Santos's taxonomy (16 sectors, 5 revenue models) was implemented exactly as "
    "provided, using Claude Haiku to classify all 4,001 companies. Work regime classification "
    "(9 categories) is planned for next semester using the Claude Batch API.",
    bold_prefix="Taxonomy implementation:"
)
bullet(doc,
    " Three additional data sources were integrated: SEC EDGAR (financial data), BLS "
    "(employment data), and Stack Overflow Developer Survey (technology adoption trends). "
    "The original plan (Scrapy + PostgreSQL) was replaced with Playwright + SQLite, which "
    "better suited the solo, iterative nature of the project.",
    bold_prefix="Technical rebuild:"
)

# ── Milestone-Specific: Contributions ────────────────────────────────────────
heading(doc, "Milestone-Specific: Individual Contributions")
body(doc, "This is a solo project. All contributions were made by Sunghun Kim:")
bullet(doc, " Custom Playwright scraper for 19 builtin.com tech hubs (7,601 raw entries)", bold_prefix="scrapers/builtin_scraper.py:")
bullet(doc, " CSV-to-SQLite loader and 4-step cleaning pipeline (deduplication, normalization, state extraction)", bold_prefix="pipeline/load_sqlite.py, clean_sqlite.py:")
bullet(doc, " Claude Haiku integration classifying 4,001 companies using Prof. Santos's taxonomy", bold_prefix="pipeline/classify_companies.py:")
bullet(doc, " SEC CIK matching + financial data retrieval for 2,937 public companies (15,952 records)", bold_prefix="pipeline/enrich_sec.py:")
bullet(doc, " BLS sector-level employment and wage data collection (11 NAICS categories)", bold_prefix="pipeline/enrich_bls.py:")
bullet(doc, " Stack Overflow survey processing across 9 years with professional developer filter", bold_prefix="pipeline/analyze_so_survey.py:")
bullet(doc, " 6-section interactive Streamlit dashboard with Plotly, deployed on Streamlit Cloud", bold_prefix="dashboard.py:")

# ── Milestone-Specific: Presentation Outline ─────────────────────────────────
heading(doc, "Milestone-Specific: Status Presentation Outline")
body(doc, "Early December Status Presentation — estimated 12–15 minutes")

heading(doc, "I. Opening (1 min)", level=2)
bullet(doc, 'Title: "Mapping the U.S. Digital Economy: Where Should the Next Investment Dollar Go?"')
bullet(doc, 'Hook: "4,001 companies. 4 data sources. One question."')

heading(doc, "II. Problem & Motivation (2 min)", level=2)
bullet(doc, "The gap: no public, queryable map of the U.S. digital economy at the sector level")
bullet(doc, "Why it matters: investors need signal, not noise")
bullet(doc, "The three-axis framework: Learning → Inventing → Investing")

heading(doc, "III. What I Built — Live Demo (3 min)", level=2)
bullet(doc, "Data pipeline overview: builtin.com → LLM classification → SEC / BLS / SO Survey")
bullet(doc, "Database model: 10 tables, 4 data layers")
bullet(doc, "Dashboard walkthrough: 30-second tour of all 6 sections")
bullet(doc, "Interactive Opportunity Score: adjust weights, watch rankings update live")

heading(doc, "IV. Key Findings (4 min)", level=2)
bullet(doc, "AI Foundation Models: highest growth (386%) + highest R&D intensity (~200%)")
bullet(doc, "Developer trend: Claude adoption +35pp in one year (8% → 43%)")
bullet(doc, "Geography: CA + TX + NY + MA = 75% of U.S. tech companies")
bullet(doc, "Opportunity Score top 3 sectors: evidence from all three data layers")

heading(doc, "V. Methodology & Limitations (2 min)", level=2)
bullet(doc, "SEC coverage: public companies only (2,937 / 4,001 = 73%)")
bullet(doc, "LLM classification: how it works, spot-check validation")
bullet(doc, "Opportunity Score: formula, weight sensitivity, what it can and cannot predict")

heading(doc, "VI. Next Steps & Q&A (2 min)", level=2)
bullet(doc, "Work regime classification (9 categories, Claude Batch API — next semester)")
bullet(doc, "Backtesting: validate Opportunity Score against historical data")
bullet(doc, "Questions")

# ── References ───────────────────────────────────────────────────────────────
heading(doc, "References")
refs = [
    "[1] Bureau of Economic Analysis. (2024). Defining and Measuring the Digital Economy. "
    "U.S. Department of Commerce. https://www.bea.gov/data/special-topics/digital-economy",
    "[2] Stack Overflow. (2017–2025). Stack Overflow Annual Developer Survey. "
    "https://survey.stackoverflow.co",
    "[3] U.S. Securities and Exchange Commission. (2024). EDGAR Full-Text Search and "
    "Company Facts API. https://efts.sec.gov / https://data.sec.gov",
    "[4] U.S. Bureau of Labor Statistics. (2025). Quarterly Census of Employment and Wages "
    "(QCEW). U.S. Department of Labor. https://www.bls.gov/cew/",
    "[5] Built In. (2024). Tech companies and startups directory. https://builtin.com/companies",
]
for ref in refs:
    p = doc.add_paragraph(ref)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Inches(0.3)

# ── Save ─────────────────────────────────────────────────────────────────────
out = "SunghunKim_Milestone2.docx"
doc.save(out)
print(f"Saved: {out}")
